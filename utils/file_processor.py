"""
Intelligent File Processing System

This module demonstrates several crucial software engineering patterns and concepts
that you'll encounter in professional development:

1. Strategy Pattern: Different processing strategies for different file types
2. Factory Pattern: Creating appropriate processors based on file characteristics
3. Chain of Responsibility: Multiple processing steps that can be combined
4. Error Handling: Graceful degradation when files can't be processed
5. Data Transformation: Converting various formats into AI-readable text

Think of this module as a universal translator that can read almost any file format
and explain its contents to our AI assistant. Just like how a skilled librarian can
quickly understand and summarize different types of documents, this system adapts
its approach based on what kind of file it encounters.
"""

import io
import os
import logging
from abc import ABC, abstractmethod
from typing import Dict, Any, Optional, List, Union, BinaryIO
from pathlib import Path
import mimetypes

# External libraries for different file formats
try:
    import PyPDF2
    import pdfplumber  # Better PDF text extraction
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False

try:
    from PIL import Image
    import pytesseract  # OCR for extracting text from images
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

import json
import csv
import base64
from datetime import datetime

from config.settings import AppConfig

logger = logging.getLogger(__name__)

class FileProcessingError(Exception):
    """
    Custom exception for file processing errors.
    
    Creating custom exceptions is a best practice because it allows us to handle
    specific types of errors differently. For example, we might retry network
    errors but not file format errors.
    """
    pass

class BaseFileProcessor(ABC):
    """
    Abstract base class for all file processors.
    
    This demonstrates the Template Method pattern - we define a common interface
    that all file processors must implement, ensuring consistency across different
    file types. Think of this as a contract that says "every file processor must
    be able to do these basic things."
    """
    
    @abstractmethod
    def can_process(self, file_type: str, file_name: str) -> bool:
        """
        Determine if this processor can handle the given file type.
        
        This method allows our system to automatically choose the right processor
        for each file, similar to how your brain automatically chooses different
        reading strategies for a novel versus a technical manual.
        """
        pass
    
    @abstractmethod
    def process(self, file_content: bytes, file_name: str, file_type: str) -> Dict[str, Any]:
        """
        Process the file and extract meaningful information.
        
        The return dictionary provides a standardized format that our AI can
        understand, regardless of the original file format. This is like having
        a universal summary format for all documents.
        """
        pass
    
    def get_file_metadata(self, file_name: str, file_size: int) -> Dict[str, Any]:
        """
        Extract basic metadata that's useful for all file types.
        
        Metadata helps users and the AI understand context about the file
        before diving into its contents.
        """
        return {
            'name': file_name,
            'size_bytes': file_size,
            'size_human': self._format_file_size(file_size),
            'extension': Path(file_name).suffix.lower(),
            'processed_at': datetime.now().isoformat()
        }
    
    def _format_file_size(self, size_bytes: int) -> str:
        """Convert bytes to human-readable format."""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"

class TextFileProcessor(BaseFileProcessor):
    """
    Processor for plain text files.
    
    Text files are the simplest case, but even here we need to handle different
    encodings and potentially large files. This processor demonstrates how to
    handle the common case efficiently while preparing for edge cases.
    """
    
    def can_process(self, file_type: str, file_name: str) -> bool:
        text_types = ['text/plain', 'text/csv', 'text/markdown', 'application/json']
        text_extensions = ['.txt', '.md', '.csv', '.json', '.log', '.py', '.js', '.html', '.css']
        
        return (file_type in text_types or 
                any(file_name.lower().endswith(ext) for ext in text_extensions))
    
    def process(self, file_content: bytes, file_name: str, file_type: str) -> Dict[str, Any]:
        """
        Process text files with intelligent encoding detection.
        
        Text encoding is one of those "invisible" technical challenges that can
        cause major headaches. This method tries different encodings gracefully,
        like a patient translator trying different languages until one works.
        """
        try:
            # Try different encodings in order of likelihood
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            text_content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    text_content = file_content.decode(encoding)
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                raise FileProcessingError("Could not decode text file with any supported encoding")
            
            # Analyze the text content
            lines = text_content.split('\n')
            word_count = len(text_content.split())
            char_count = len(text_content)
            
            # Extract a meaningful preview
            preview = self._create_text_preview(text_content)
            
            # Try to detect the specific text format for better processing
            text_format = self._detect_text_format(text_content, file_name)
            
            return {
                'type': 'text',
                'format': text_format,
                'content': text_content,
                'preview': preview,
                'metadata': self.get_file_metadata(file_name, len(file_content)),
                'statistics': {
                    'line_count': len(lines),
                    'word_count': word_count,
                    'character_count': char_count,
                    'encoding_used': used_encoding
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing text file {file_name}: {str(e)}")
            raise FileProcessingError(f"Failed to process text file: {str(e)}")
    
    def _create_text_preview(self, text: str, max_chars: int = 500) -> str:
        """
        Create an intelligent preview of text content.
        
        Instead of just truncating text arbitrarily, we try to break at natural
        boundaries like sentences or paragraphs to create more readable previews.
        """
        if len(text) <= max_chars:
            return text
        
        # Try to break at sentence boundaries
        preview = text[:max_chars]
        last_sentence = max(preview.rfind('.'), preview.rfind('!'), preview.rfind('?'))
        
        if last_sentence > max_chars * 0.7:  # If we found a good sentence break
            preview = preview[:last_sentence + 1]
        else:
            # Fall back to word boundaries
            preview = preview.rsplit(' ', 1)[0] + '...'
        
        return preview
    
    def _detect_text_format(self, content: str, filename: str) -> str:
        """
        Try to determine the specific format of text content.
        
        This helps our AI understand how to interpret the content. For example,
        JSON should be parsed differently than free-form text.
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.json'):
            return 'json'
        elif filename_lower.endswith('.csv'):
            return 'csv'
        elif filename_lower.endswith('.md'):
            return 'markdown'
        elif filename_lower.endswith(('.py', '.js', '.html', '.css')):
            return 'code'
        elif content.strip().startswith('{') and content.strip().endswith('}'):
            return 'json'
        elif ',' in content and '\n' in content:
            # Simple heuristic for CSV detection
            lines = content.split('\n')[:5]
            if all(',' in line for line in lines if line.strip()):
                return 'csv'
        
        return 'plain_text'

class PDFProcessor(BaseFileProcessor):
    """
    Processor for PDF documents.
    
    PDFs are complex because they can contain text, images, tables, and complex
    layouts. This processor demonstrates how to extract meaningful content from
    structured documents while handling various edge cases.
    """
    
    def can_process(self, file_type: str, file_name: str) -> bool:
        return (PDF_AVAILABLE and 
                (file_type == 'application/pdf' or file_name.lower().endswith('.pdf')))
    
    def process(self, file_content: bytes, file_name: str, file_type: str) -> Dict[str, Any]:
        """
        Extract text and metadata from PDF files.
        
        We use multiple extraction methods because PDFs can be tricky - some have
        selectable text, others are scanned images, and some have complex layouts
        that require special handling.
        """
        if not PDF_AVAILABLE:
            raise FileProcessingError("PDF processing libraries not available")
        
        try:
            # Create a file-like object from bytes
            pdf_file = io.BytesIO(file_content)
            
            # Try pdfplumber first for better text extraction
            text_content = self._extract_with_pdfplumber(pdf_file)
            
            # Fallback to PyPDF2 if pdfplumber fails
            if not text_content.strip():
                pdf_file.seek(0)  # Reset file pointer
                text_content = self._extract_with_pypdf2(pdf_file)
            
            # Get PDF metadata
            pdf_file.seek(0)
            metadata = self._extract_pdf_metadata(pdf_file)
            
            if not text_content.strip():
                text_content = "[PDF appears to contain mainly images or has no extractable text]"
            
            # Create a preview
            preview = self._create_text_preview(text_content)
            
            return {
                'type': 'pdf',
                'content': text_content,
                'preview': preview,
                'metadata': {
                    **self.get_file_metadata(file_name, len(file_content)),
                    **metadata
                },
                'statistics': {
                    'character_count': len(text_content),
                    'word_count': len(text_content.split()),
                    'extractable_text': bool(text_content.strip())
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_name}: {str(e)}")
            raise FileProcessingError(f"Failed to process PDF: {str(e)}")
    
    def _extract_with_pdfplumber(self, pdf_file: BinaryIO) -> str:
        """
        Extract text using pdfplumber for better layout preservation.
        
        pdfplumber is better at maintaining the structure and layout of text,
        which is important for understanding tables and formatted content.
        """
        try:
            import pdfplumber
            text_parts = []
            
            with pdfplumber.open(pdf_file) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
            
            return '\n\n'.join(text_parts)
            
        except ImportError:
            return ""  # Fall back to PyPDF2
        except Exception as e:
            logger.warning(f"pdfplumber extraction failed: {str(e)}")
            return ""
    
    def _extract_with_pypdf2(self, pdf_file: BinaryIO) -> str:
        """
        Fallback text extraction using PyPDF2.
        
        PyPDF2 is more basic but more widely compatible. It's our safety net
        when more sophisticated extraction methods fail.
        """
        try:
            text_parts = []
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            logger.warning(f"PyPDF2 extraction failed: {str(e)}")
            return ""
    
    def _extract_pdf_metadata(self, pdf_file: BinaryIO) -> Dict[str, Any]:
        """
        Extract metadata from PDF for additional context.
        
        PDF metadata can include author, creation date, subject, and other
        information that helps understand the document's context and purpose.
        """
        try:
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            metadata = pdf_reader.metadata
            
            result = {
                'page_count': len(pdf_reader.pages),
                'title': getattr(metadata, 'title', None),
                'author': getattr(metadata, 'author', None),
                'subject': getattr(metadata, 'subject', None),
                'creator': getattr(metadata, 'creator', None),
                'creation_date': getattr(metadata, 'creation_date', None),
                'modification_date': getattr(metadata, 'modification_date', None)
            }
            
            # Clean up None values and convert dates to strings
            return {k: str(v) if v is not None else None for k, v in result.items()}
            
        except Exception as e:
            logger.warning(f"Could not extract PDF metadata: {str(e)}")
            return {'page_count': 'unknown'}

class DocxProcessor(BaseFileProcessor):
    """
    Processor for Microsoft Word documents.
    
    Word documents can contain rich formatting, tables, images, and embedded
    objects. This processor focuses on extracting the text content while
    preserving structural information where possible.
    """
    
    def can_process(self, file_type: str, file_name: str) -> bool:
        return (DOCX_AVAILABLE and 
                (file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or
                 file_name.lower().endswith('.docx')))
    
    def process(self, file_content: bytes, file_name: str, file_type: str) -> Dict[str, Any]:
        """
        Extract text and structure from Word documents.
        
        Word documents have rich structure - paragraphs, headings, tables, etc.
        We try to preserve this structure in our text extraction to help the AI
        understand the document's organization.
        """
        if not DOCX_AVAILABLE:
            raise FileProcessingError("DOCX processing library not available")
        
        try:
            # Create document from bytes
            doc_file = io.BytesIO(file_content)
            doc = DocxDocument(doc_file)
            
            # Extract text with structure preservation
            text_parts = []
            paragraph_count = 0
            table_count = 0
            
            for element in doc.element.body:
                if element.tag.endswith('p'):  # Paragraph
                    para = doc.paragraphs[paragraph_count]
                    if para.text.strip():
                        # Check if it's a heading
                        if para.style.name.startswith('Heading'):
                            text_parts.append(f"\n{para.style.name}: {para.text}\n")
                        else:
                            text_parts.append(para.text)
                    paragraph_count += 1
                    
                elif element.tag.endswith('tbl'):  # Table
                    if table_count < len(doc.tables):
                        table = doc.tables[table_count]
                        text_parts.append(self._extract_table_text(table))
                    table_count += 1
            
            full_text = '\n'.join(text_parts)
            
            # Get document properties
            properties = self._extract_docx_properties(doc)
            
            return {
                'type': 'docx',
                'content': full_text,
                'preview': self._create_text_preview(full_text),
                'metadata': {
                    **self.get_file_metadata(file_name, len(file_content)),
                    **properties
                },
                'statistics': {
                    'paragraph_count': len(doc.paragraphs),
                    'table_count': len(doc.tables),
                    'character_count': len(full_text),
                    'word_count': len(full_text.split())
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_name}: {str(e)}")
            raise FileProcessingError(f"Failed to process Word document: {str(e)}")
    
    def _extract_table_text(self, table) -> str:
        """
        Extract text from Word tables in a readable format.
        
        Tables need special handling to maintain readability. We format them
        as text in a way that preserves the relationship between cells.
        """
        table_text = ["\n[TABLE]"]
        
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_cells.append(cell_text)
            table_text.append(" | ".join(row_cells))
        
        table_text.append("[/TABLE]\n")
        return '\n'.join(table_text)
    
    def _extract_docx_properties(self, doc) -> Dict[str, Any]:
        """Extract document properties and metadata."""
        try:
            props = doc.core_properties
            return {
                'title': props.title,
                'author': props.author,
                'subject': props.subject,
                'keywords': props.keywords,
                'created': props.created.isoformat() if props.created else None,
                'modified': props.modified.isoformat() if props.modified else None,
                'last_modified_by': props.last_modified_by
            }
        except Exception as e:
            logger.warning(f"Could not extract document properties: {str(e)}")
            return {}

class CSVProcessor(BaseFileProcessor):
    """
    Processor for CSV (Comma-Separated Values) files.
    
    CSV files are deceptively simple but can have many variations in format,
    encoding, and structure. This processor demonstrates how to handle tabular
    data intelligently and make it accessible to AI analysis.
    """
    
    def can_process(self, file_type: str, file_name: str) -> bool:
        return (PANDAS_AVAILABLE and 
                (file_type == 'text/csv' or file_name.lower().endswith('.csv')))
    
    def process(self, file_content: bytes, file_name: str, file_type: str) -> Dict[str, Any]:
        """
        Process CSV files with intelligent structure detection.
        
        CSV files can vary widely in format - different delimiters, quote characters,
        encoding, and structure. This method tries to automatically detect these
        characteristics and present the data in a useful format.
        """
        if not PANDAS_AVAILABLE:
            raise FileProcessingError("Pandas library not available for CSV processing")
        
        try:
            # Convert bytes to string with encoding detection
            text_content = self._decode_csv_content(file_content)
            
            # Detect CSV dialect (delimiter, quote character, etc.)
            dialect = self._detect_csv_dialect(text_content)
            
            # Read CSV with pandas for intelligent handling
            csv_file = io.StringIO(text_content)
            df = pd.read_csv(csv_file, dialect=dialect, low_memory=False)
            
            # Generate summary statistics
            summary = self._generate_csv_summary(df)
            
            # Create a preview of the data
            preview = self._create_csv_preview(df)
            
            # Convert to text format for AI processing
            csv_text = self._convert_csv_to_text(df)
            
            return {
                'type': 'csv',
                'content': csv_text,
                'preview': preview,
                'metadata': {
                    **self.get_file_metadata(file_name, len(file_content)),
                    'columns': list(df.columns),
                    'shape': df.shape,
                    'dialect': {
                        'delimiter': dialect.delimiter,
                        'quotechar': dialect.quotechar,
                        'lineterminator': repr(dialect.lineterminator)
                    }
                },
                'statistics': summary
            }
            
        except Exception as e:
            logger.error(f"Error processing CSV {file_name}: {str(e)}")
            raise FileProcessingError(f"Failed to process CSV file: {str(e)}")
    
    def _decode_csv_content(self, file_content: bytes) -> str:
        """Decode CSV content with encoding detection."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'utf-16']
        
        for encoding in encodings:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        raise FileProcessingError("Could not decode CSV file with any supported encoding")
    
    def _detect_csv_dialect(self, content: str):
        """
        Detect CSV format characteristics automatically.
        
        This uses Python's csv.Sniffer to automatically detect things like
        whether commas or semicolons are used as delimiters, which helps us
        parse files correctly regardless of their specific format.
        """
        try:
            # Use a sample to detect dialect
            sample = content[:1024]
            sniffer = csv.Sniffer()
            dialect = sniffer.sniff(sample, delimiters=',;\t|')
            return dialect
        except Exception:
            # Fall back to default comma-separated format
            class DefaultDialect:
                delimiter = ','
                quotechar = '"'
                lineterminator = '\n'
            return DefaultDialect()
    
    def _generate_csv_summary(self, df: 'pd.DataFrame') -> Dict[str, Any]:
        """
        Generate statistical summary of CSV data.
        
        This creates insights about the data that help users and the AI understand
        what kind of information the CSV contains before diving into specific analysis.
        """
        numeric_columns = df.select_dtypes(include=['number']).columns.tolist()
        text_columns = df.select_dtypes(include=['object']).columns.tolist()
        
        summary = {
            'row_count': len(df),
            'column_count': len(df.columns),
            'numeric_columns': len(numeric_columns),
            'text_columns': len(text_columns),
            'missing_values': df.isnull().sum().sum(),
            'duplicate_rows': df.duplicated().sum()
        }
        
        # Add statistics for numeric columns
        if numeric_columns:
            summary['numeric_stats'] = {}
            for col in numeric_columns[:5]:  # Limit to first 5 numeric columns
                summary['numeric_stats'][col] = {
                    'mean': float(df[col].mean()) if not df[col].isnull().all() else None,
                    'median': float(df[col].median()) if not df[col].isnull().all() else None,
                    'min': float(df[col].min()) if not df[col].isnull().all() else None,
                    'max': float(df[col].max()) if not df[col].isnull().all() else None
                }
        
        return summary
    
    def _create_csv_preview(self, df: 'pd.DataFrame') -> str:
        """Create a readable preview of CSV data."""
        preview_rows = min(10, len(df))
        preview_df = df.head(preview_rows)
        
        # Convert to string with nice formatting
        return f"Preview of first {preview_rows} rows:\n\n{preview_df.to_string(index=False)}"
    
    def _convert_csv_to_text(self, df: 'pd.DataFrame') -> str:
        """
        Convert CSV data to a text format that AI can easily understand.
        
        We create a narrative description of the data rather than just dumping
        the raw CSV, which helps the AI provide better analysis and insights.
        """
        text_parts = [
            f"CSV Data Summary:",
            f"- {len(df)} rows and {len(df.columns)} columns",
            f"- Columns: {', '.join(df.columns.tolist())}",
            "",
            "Sample data:"
        ]
        
        # Include a reasonable sample of the data
        sample_size = min(20, len(df))
        sample_df = df.head(sample_size)
        
        text_parts.append(sample_df.to_string(index=False))
        
        if len(df) > sample_size:
            text_parts.append(f"\n... and {len(df) - sample_size} more rows")
        
        return '\n'.join(text_parts)

class ImageProcessor(BaseFileProcessor):
    """
    Processor for image files with optional OCR capabilities.
    
    Images can contain text (which we can extract with OCR) or be purely visual.
    This processor demonstrates how to handle both cases and provide useful
    information about image content.
    """
    
    def can_process(self, file_type: str, file_name: str) -> bool:
        image_types = ['image/jpeg', 'image/png', 'image/gif', 'image/bmp', 'image/webp']
        image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp']
        
        return (file_type in image_types or 
                any(file_name.lower().endswith(ext) for ext in image_extensions))
    
    def process(self, file_content: bytes, file_name: str, file_type: str) -> Dict[str, Any]:
        """
        Process image files with optional text extraction.
        
        Images are processed differently than text files - we can extract metadata
        about the image itself (dimensions, format, etc.) and optionally try to
        extract any text that appears in the image using OCR.
        """
        try:
            # Load image and get basic information
            image = Image.open(io.BytesIO(file_content))
            
            # Extract image metadata
            image_info = {
                'format': image.format,
                'mode': image.mode,
                'size': image.size,
                'width': image.width,
                'height': image.height
            }
            
            # Try to extract text using OCR if available
            extracted_text = ""
            if OCR_AVAILABLE:
                try:
                    extracted_text = pytesseract.image_to_string(image)
                    extracted_text = extracted_text.strip()
                except Exception as e:
                    logger.warning(f"OCR failed for {file_name}: {str(e)}")
            
            # Convert image to base64 for potential display
            image_b64 = base64.b64encode(file_content).decode('utf-8')
            
            # Create content description
            content_parts = [f"Image file: {file_name}"]
            content_parts.append(f"Dimensions: {image.width}x{image.height} pixels")
            content_parts.append(f"Format: {image.format}")
            
            if extracted_text:
                content_parts.append(f"\nText extracted from image:\n{extracted_text}")
            else:
                content_parts.append("\nNo text detected in image or OCR not available.")
            
            content = '\n'.join(content_parts)
            
            return {
                'type': 'image',
                'content': content,
                'preview': f"Image: {file_name} ({image.width}x{image.height})",
                'extracted_text': extracted_text,
                'image_base64': image_b64,
                'metadata': {
                    **self.get_file_metadata(file_name, len(file_content)),
                    **image_info
                },
                'statistics': {
                    'has_text': bool(extracted_text),
                    'text_length': len(extracted_text) if extracted_text else 0,
                    'file_size_bytes': len(file_content)
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing image {file_name}: {str(e)}")
            raise FileProcessingError(f"Failed to process image: {str(e)}")

class FileProcessor:
    """
    Main file processing coordinator that manages different processor types.
    
    This class implements the Strategy pattern - it maintains a collection of
    different processing strategies and chooses the appropriate one based on
    the file type. Think of it as a smart dispatcher that knows which specialist
    to call for each type of document.
    """
    
    def __init__(self):
        self.config = AppConfig()
        
        # Initialize all available processors
        self.processors = [
            TextFileProcessor(),
            PDFProcessor(),
            DocxProcessor(),
            CSVProcessor(),
            ImageProcessor()
        ]
        
        logger.info(f"FileProcessor initialized with {len(self.processors)} processors")
    
    def process_file(self, uploaded_file) -> str:
        """
        Process an uploaded file and return extracted content.
        
        This is the main entry point for file processing. It demonstrates how
        to coordinate between different specialized processors while providing
        a simple, consistent interface to the rest of the application.
        """
        try:
            # Read file content
            file_content = uploaded_file.read()
            file_name = uploaded_file.name
            file_type = uploaded_file.type or self._guess_file_type(file_name)
            
            # Validate file size
            if len(file_content) > self.config.get_max_file_size_bytes():
                raise FileProcessingError(
                    f"File too large. Maximum size is {self.config.ui.max_file_size_mb}MB"
                )
            
            # Find appropriate processor
            processor = self._find_processor(file_type, file_name)
            
            if not processor:
                return self._create_fallback_content(file_name, file_type, len(file_content))
            
            # Process the file
            result = processor.process(file_content, file_name, file_type)
            
            # Return the extracted content
            return result.get('content', 'No content could be extracted from this file.')
            
        except FileProcessingError:
            raise  # Re-raise our custom exceptions
        except Exception as e:
            logger.error(f"Unexpected error processing file {uploaded_file.name}: {str(e)}")
            raise FileProcessingError(f"Unexpected error processing file: {str(e)}")
    
    def _find_processor(self, file_type: str, file_name: str) -> Optional[BaseFileProcessor]:
        """
        Find the appropriate processor for a given file type.
        
        This method demonstrates how the Strategy pattern works in practice -
        we ask each processor if it can handle the file, and use the first one
        that says yes. The order matters here, with more specific processors
        coming before more general ones.
        """
        for processor in self.processors:
            if processor.can_process(file_type, file_name):
                logger.info(f"Using {processor.__class__.__name__} for {file_name}")
                return processor
        
        logger.warning(f"No processor found for file type: {file_type}, name: {file_name}")
        return None
    
    def _guess_file_type(self, file_name: str) -> str:
        """
        Guess file type from filename when MIME type is not available.
        
        Sometimes the browser doesn't provide a MIME type, so we fall back to
        guessing based on the file extension. This is less reliable but better
        than nothing.
        """
        mime_type, _ = mimetypes.guess_type(file_name)
        return mime_type or 'application/octet-stream'
    
    def _create_fallback_content(self, file_name: str, file_type: str, file_size: int) -> str:
        """
        Create basic content description when we can't process a file.
        
        Even when we can't extract content from a file, we can still provide
        useful information about what the file is. This ensures users always
        get some kind of meaningful response.
        """
        size_mb = file_size / (1024 * 1024)
        
        return f"""
        File: {file_name}
        Type: {file_type}
        Size: {size_mb:.2f} MB
        
        This file type is not currently supported for content extraction, but I can see
        basic information about it. You can ask me questions about what this type of
        file typically contains or how to work with it.
        """
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """
        Get information about supported file formats.
        
        This provides users with clear information about what file types the
        application can handle, helping them understand the system's capabilities.
        """
        formats = {
            'Text Files': ['.txt', '.md', '.csv', '.json', '.log', '.py', '.js', '.html', '.css'],
            'Documents': [],
            'Images': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp'],
            'Data Files': []
        }
        
        if PDF_AVAILABLE:
            formats['Documents'].append('.pdf')
        
        if DOCX_AVAILABLE:
            formats['Documents'].append('.docx')
        
        if PANDAS_AVAILABLE:
            formats['Data Files'].append('.csv')
        
        return formats
    
    def get_processing_capabilities(self) -> Dict[str, str]:
        """
        Get information about what the processor can extract from different file types.
        
        This helps users understand not just what files are supported, but what
        kind of information can be extracted from each type.
        """
        capabilities = {
            'Text Files': 'Full text content, encoding detection, format analysis',
            'Images': 'Image metadata, dimensions, format information',
        }
        
        if PDF_AVAILABLE:
            capabilities['PDF Documents'] = 'Text extraction, metadata, page count, structure preservation'
        
        if DOCX_AVAILABLE:
            capabilities['Word Documents'] = 'Text content, document structure, tables, metadata'
        
        if PANDAS_AVAILABLE:
            capabilities['CSV Files'] = 'Data analysis, statistics, column information, preview'
        
        if OCR_AVAILABLE:
            capabilities['Images'] += ', text extraction (OCR)'
        
        return capabilities